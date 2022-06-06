from docx import Document
from easygui import fileopenbox
from fileinput import FileInput


def text_replacer(search, change):

    extensions = ['docx', 'txt']

    # Select file
    # Check extension for validity
    try:
        file_path = fileopenbox()

        file_path_checker = file_path.split('.')

    except AttributeError:
        print('Program exited...')

    else:
        if file_path_checker[-1] not in extensions:

            print('Sorry, this program can\'t access that file type.')

        elif file_path_checker[-1] == 'txt':

            # Open, read, overwrite keywords, save.
            with FileInput(file_path, inplace=True) as file:
                for line in file:
                    print(line.replace(search, change))

            print(
                f'All instances of {search} have been replaced with {change}!')

        else:
            print("Still working on that feature!")
            doc = Document(file_path)
            doc.save('UpdatedCoverLetter.docx')
            doc = Document('UpdatedCoverLetter.docx')

            for p in doc.paragraphs:
                if search in p.text:
                    inline = p.runs
                    # Loop added to work with runs (strings with same style)
                    for i in range(len(inline)):
                        if search in inline[i].text:
                            text = inline[i].text.replace(
                                search, change)
                            print(text)
                            inline[i].text = text

            doc.save('UpdatedCoverLetter.docx')


text_replacer('company', 'wonko')
